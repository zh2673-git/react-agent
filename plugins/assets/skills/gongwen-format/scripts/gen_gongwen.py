#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""按 GB/T 9704-2012 生成党政机关公文 docx（确定性排版）。

用法:
    python scripts/gen_gongwen.py doc.json 输出.docx

doc.json 结构（示意，字段均可选）:
{
  "doc_type": "下行文 | 上行文",              // 默认下行文
  "top_fields": {
    "fenhao": "000001",                        // 份号（6 位，可选，版心左上角第一行）
    "miji": "机密★5年",                        // 密级（可选，3 号黑体，左上角第二行）
    "jinji": "特急",                           // 紧急程度（可选，3 号黑体）
    "signer": "张三"                           // 签发人姓名（上行文用）
  },
  "party": "×××市人民政府",                   // 发文机关（红头，红色，居中）
  "year": "2026",                              // 发文字号年份（可选；缺省从 add 推断）
  "no": "5",                                   // 发文字号顺序号（可选）
  "add": "〔2026〕5号",                        // 发文字号完整字符串（优先于 year/no）
  "title": "关于××的通知",
  "main_send": "各区县人民政府、市政府各部门：",
  "content": [                                 // 正文段落；type 见下
    {"type": "h1", "text": "一、总体要求"},
    {"type": "h2", "text": "（一）基本原则"},
    {"type": "h3", "text": "1. 市场主导"},
    {"type": "h4", "text": "（1）任务分工"},
    {"type": "para", "text": "正文……"}
  ],
  "attachment_notes": "附件：1.×××措施",        // 附件说明（可选）
  "signing": {                                  // 落款
    "org": "×××市人民政府",                    // 署名机关
    "date": "2026年9月7日",                    // 成文日期（月日不编虚位）
    "stamp": true                               // true=盖章式；false=不盖章式
  },
  "note": "（此件公开发布）",                    // 附注（可选，3 号仿宋，圆括号）
  "attachments": [                              // 附件另面（可选）
    {"index": 1, "title": "附件标题", "paragraphs": [同 content]}
  ],
  "banji": {                                    // 版记（4 号仿宋，位于最后一面）
    "send": "抄送：×××。",
    "print_org": "×××市人民政府办公室",        // 印发机关
    "print_date": "2026年9月7日印发"            // 印发日期（含"印发"）
  }
}

字体自动降级：方正小标宋→华文中宋(STZhongsong)→宋体加粗；仿宋→仿宋→华文仿宋。
发生降级时终端打印提示，可到 Word 中再人工核对。

注：标准中"红头上边缘至版心上边缘 35mm""分隔线在发文字号下 4mm""版记末线与版心
下边缘重合"等绝对定位在 docx 引擎中难以像素级表达，脚本按等效的结构顺序与行数
近似排版；要求精确定位的场景请在 Word 中按 references 微调。
"""
import json
import os
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

PAGE_W, PAGE_H = Cm(21.0), Cm(29.7)
TOP, BOTTOM, LEFT, RIGHT = Cm(3.7), Cm(3.5), Cm(2.8), Cm(2.6)
RED = RGBColor(0xFF, 0x00, 0x00)

# 逻辑字体 -> 首选/降级链（Windows 字体名）
FONT_CHAINS = {
    "xiaobiaosong": ["方正小标宋简体", "华文中宋", "STZhongsong", "宋体"],
    "fangsong": ["仿宋", "仿宋_GB2312", "宋体"],
    "songti": ["宋体"],
    "heiti": ["黑体"],
    "kaiti": ["楷体", "KaiTi", "宋体"],
}


def _pick(chain):
    """返回本机已安装的第一个字体名；探测失败返回链首（交由 Word 渲染）。"""
    try:
        import winreg
        key = winreg.OpenKey(winreg.HKEY_LOCAL_MACHINE,
                             r"SOFTWARE\Microsoft\Windows NT\CurrentVersion\Fonts")
        installed = []
        i = 0
        while True:
            try:
                installed.append(winreg.EnumValue(key, i)[0])
                i += 1
            except OSError:
                break
        winreg.CloseKey(key)
    except Exception:
        return chain[0]
    for f in chain:
        for name in installed:
            if f == "方正小标宋简体" and "方正小标宋" in name:
                return f
            if f == "华文中宋" and "华文中宋" in name:
                return "华文中宋"
            if f == "STZhongsong" and "华文中宋" in name:
                return "华文中宋"
            if f in ("仿宋", "仿宋_GB2312") and f in name:
                return f
            if f == "宋体" and name.startswith("宋体") and "新宋体" not in name:
                return "宋体"
            if f == "黑体" and name.startswith("黑体"):
                return "黑体"
            if f in ("楷体", "KaiTi") and (name.startswith("楷体") or "KaiTi" in name):
                return "楷体" if name.startswith("楷体") else f
    return chain[0]


def _set_font(run, chain, size, bold=False, color=None, east="", ascii_font="Times New Roman"):
    run.font.size = Pt(size)
    run.font.name = ascii_font
    if bold:
        run.font.bold = True
    if color is not None:
        run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    font = _pick(FONT_CHAINS[chain])
    rFonts.set(qn("w:ascii"), ascii_font)
    rFonts.set(qn("w:hAnsi"), ascii_font)
    rFonts.set(qn("w:eastAsia"), east or font)
    return font


def _fixed_line(p, pt):
    pPr = p._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:line"), str(int(pt * 20)))
    spacing.set(qn("w:lineRule"), "exact")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")


def _indent(p, chars=2):
    """首行缩进 N 字符（用 firstLineChars，等宽可靠）。"""
    pPr = p._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    ind.set(qn("w:firstLineChars"), str(chars * 100))
    ind.set(qn("w:firstLine"), str(int(16 * 20 * chars)))  # 3 号字 16pt 近似


def _right_indent(p, chars):
    """右缩进 N 字符（用于右空 X 字）。"""
    pPr = p._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    ind.set(qn("w:rightChars"), str(chars * 100))


def _bottom_border(p, size_eighths, color="000000"):
    """段落加下边框（模拟分隔线）。size 单位 1/8 pt：8=1磅粗，6=0.75磅细。"""
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size_eighths))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _para(doc, text="", chain="fangsong", size=16, align=None, first_indent=0,
          right_chars=0, bold=False, color=None, fixed=28, border=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if fixed:
        _fixed_line(p, fixed)
    if first_indent:
        _indent(p, first_indent)
    if right_chars:
        _right_indent(p, right_chars)
    if border:
        _bottom_border(p, border[0], border[1])
    if text:
        r = p.add_run(text)
        _set_font(r, chain, size, bold=bold, color=color)
    return p


def add_page_number(footer_para):
    """页脚插入 — n —（域自动编号），宋体 4 号。"""
    def _r(t, is_field=False, instr=""):
        r = footer_para.add_run()
        _set_font(r, "songti", 14)
        if is_field:
            b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
            it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
            e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
            r._r.append(b); r._r.append(it); r._r.append(e)
        else:
            r.text = t
        return r
    _r("— ")
    _r("", True, " PAGE ")
    _r(" —")


def main():
    if len(sys.argv) < 3:
        print(__doc__)
        sys.exit(1)
    inp, out = sys.argv[1], sys.argv[2]
    with open(inp, encoding="utf-8") as f:
        data = json.load(f)

    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = PAGE_W, PAGE_H
    sec.top_margin, sec.bottom_margin = TOP, BOTTOM
    sec.left_margin, sec.right_margin = LEFT, RIGHT

    doc_type = data.get("doc_type", "下行文")
    top = data.get("top_fields") or {}
    signer = top.get("signer")
    party = data.get("party", "")
    add = data.get("add", "").strip()
    if not add:
        add = "〔%s〕%s号" % (data.get("year", "2026"), data.get("no", "1"))
    title = data.get("title", "")
    main_send = data.get("main_send", "")

    # ------- 版头 -------
    if top.get("fenhao"):
        _para(doc, top["fenhao"], "songti", 16, bold=True, fixed=28)
    if top.get("miji") or top.get("jinji"):
        _para(doc, " ".join(x for x in (top.get("miji"), top.get("jinji")) if x),
              "heiti", 16, bold=True, fixed=28)

    # 红头（发文机关标志，红色小标宋，居中）
    _para(doc, party, "xiaobiaosong", 36, align=WD_ALIGN_PARAGRAPH.CENTER,
          bold=True, color=RED, fixed=None)

    # 发文字号 + 签发人（同一逻辑行）
    is_up = doc_type == "上行文"
    line_p = doc.add_paragraph()
    _fixed_line(line_p, 28)
    if is_up and signer:
        # 发文字号居左空一字；"签发人：xxx" 靠右制表符对齐
        _para_in(line_p, " " + add, "songti", 16)
        _right_tab(line_p)
        r = line_p.add_run("\t签发人：")
        _set_font(r, "fangsong", 16)
        r2 = line_p.add_run(signer)
        _set_font(r2, "kaiti", 16)
    else:
        line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = line_p.add_run(add)
        _set_font(r, "songti", 16)

    # 红色分隔线（发文字号下，全宽 1 磅红）
    _para(doc, "", fixed=28, border=(8, "FF0000"))

    # ------- 主体 -------
    # 标题（分隔线下空二行：上方已有一分隔段 + 再空一行）
    _para(doc, "", fixed=28)
    _para(doc, title, "xiaobiaosong", 22, align=WD_ALIGN_PARAGRAPH.CENTER,
          bold=True, fixed=28)
    # 主送机关（标题下空一行）
    _para(doc, "", fixed=28)
    _para(doc, main_send, "fangsong", 16, fixed=28)
    # 正文
    for item in data.get("content", []):
        t = item.get("type", "para")
        txt = item.get("text", "")
        if t in ("h1", "heading1"):
            _para(doc, txt, "heiti", 16, first_indent=2, fixed=28)
        elif t in ("h2", "heading2"):
            _para(doc, txt, "kaiti", 16, first_indent=2, fixed=28)
        else:
            _para(doc, txt, "fangsong", 16, first_indent=2, fixed=28)
    # 附件说明
    if data.get("attachment_notes"):
        _para(doc, "", fixed=28)
        _para(doc, data["attachment_notes"], "fangsong", 16, first_indent=2, fixed=28)

    # 落款
    signing = data.get("signing") or {}
    org = signing.get("org", "")
    date_s = signing.get("date", "")
    stamp = signing.get("stamp", True)
    if org or date_s:
        if stamp:
            # 盖章式：署位居中，日期右空四字
            _para(doc, org, "fangsong", 16, align=WD_ALIGN_PARAGRAPH.CENTER, fixed=28)
            _para(doc, date_s, "fangsong", 16, right_chars=4, fixed=28)
        else:
            # 不盖章式：署名右空二字，日期右空四字（日期比署名右移二字）
            _para(doc, org, "fangsong", 16, right_chars=2, fixed=28)
            _para(doc, date_s, "fangsong", 16, right_chars=4, fixed=28)
    if data.get("note"):
        _para(doc, data["note"], "fangsong", 16, first_indent=2, fixed=28)

    # 附件另面
    for att in data.get("attachments", []):
        doc.add_page_break()
        _para(doc, "附件%d" % att.get("index", 1), "heiti", 16, bold=True, fixed=28)
        _para(doc, "", fixed=28)
        _para(doc, att.get("title", ""), "xiaobiaosong", 22,
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, fixed=28)
        for it in att.get("paragraphs", []):
            _para(doc, it.get("text", ""), "fangsong", 16, first_indent=2, fixed=28)

    # ------- 版记（4 号仿宋，最后一面）-------
    banji = data.get("banji") or {}
    # 首条粗线
    _para(doc, "", fixed=20, border=(8, "000000"))
    if banji.get("send"):
        # 抄送左右各空一字
        _para(doc, " " + banji["send"] + " ", "fangsong", 14, fixed=20)
    if banji.get("print_org") or banji.get("print_date"):
        row = "{}　{}".format(banji.get("print_org", ""),
                             banji.get("print_date", "")).strip()
        # 印发机关左空一字、印发日期右空一字
        _para(doc, " " + row + " ", "fangsong", 14, fixed=20)
    # 末条粗线
    _para(doc, "", fixed=20, border=(8, "000000"))

    # 页码：单页右空一字（页脚居右默认；双页左空一字需奇偶页，简化用单页）
    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _fixed_line(fp, 14)
    add_page_number(fp)

    doc.save(out)
    print("已生成:", os.path.abspath(out))
    print("提示：红头 35mm、分隔线 4mm、版纪末线贴底等绝对定位需在 Word 微调；")
    print("     若标题字体缺方正小标宋将降级为华文中宋/宋体，请注意核对。")


def _para_in(p, text, chain, size, bold=False):
    r = p.add_run(text)
    _set_font(r, chain, size, bold=bold)
    return r


def _right_tab(p):
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), "8800")  # 版心宽 156mm≈4425 twip? 实为 156mm*56.7≈8845 twip
    tabs.append(tab)
    pPr.append(tabs)


if __name__ == "__main__":
    main()
